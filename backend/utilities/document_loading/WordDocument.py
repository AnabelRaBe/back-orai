"""
This module contains the WordDocumentLoading class which inherits from DocumentLoadingBase. 
It is used to download a Word document from a given URL, convert its headings to markdown tags, 
and load the document along with its metadata into a list of SourceDocument objects.

Classes:
    WordDocumentLoading: A class to download and load Word documents.

Functions:
    __init__(self): Initializes the WordDocumentLoading object.
    _download_document(self, document_url: str): Downloads the Word document from the given URL.
    _get_opening_tag(self, heading_level: int): Returns the opening markdown tag for the given heading level.
    _get_closing_tag(self, heading_level: int): Returns the closing markdown tag for the given heading level.
    load(self, document_url: str, metadata: dict): Loads the document and its metadata into a list of SourceDocument objects.
"""

from typing import List
from io import BytesIO
from docx import Document
import requests
from .DocumentLoadingBase import DocumentLoadingBase
from ..common.SourceDocument import SourceDocument

class WordDocumentLoading(DocumentLoadingBase):
    """
    A class for loading Word documents and converting them to markdown format.

    Attributes:
        doc_headings_to_markdown_tags (dict): A dictionary mapping Word heading styles to markdown tags.

    Methods:
        _download_document(document_url: str) -> BytesIO: Downloads the document from the given URL and returns it as a BytesIO object.
        _get_opening_tag(heading_level: int) -> str: Returns the opening markdown tag for the given heading level.
        _get_closing_tag(heading_level: int) -> str: Returns the closing markdown tag for the given heading level.
        load(document_url: str, metadata: dict) -> List[SourceDocument]: Loads the Word document, converts it to markdown format, and returns a list of SourceDocument objects.
    """
  
    def __init__(self) -> None:
        super().__init__()
        self.doc_headings_to_markdown_tags = {
            'Heading 1' : 'h1',
            'Heading 2' : 'h2',
            'Heading 3' : 'h3',
            'Heading 4' : 'h4',
            'Heading 5' : 'h5',
            'Heading 6' : 'h6',
        }
    
    def _download_document(self, document_url: str) -> BytesIO:
        response = requests.get(document_url)
        file = BytesIO(response.content)
        return file
    
    def _get_opening_tag(self, heading_level: int) -> str:
        return f"<{self.doc_headings_to_markdown_tags.get(f'{heading_level}', '')}>"
    
    def _get_closing_tag(self, heading_level: int) -> str:
        return f"</{self.doc_headings_to_markdown_tags.get(f'{heading_level}', '')}>"
    
    def load(self, document_url: str, metadata: dict) -> List[SourceDocument]:
        output = ""
        document = Document(self._download_document(document_url))
        for paragraph in document.paragraphs:
            output += f"{self._get_opening_tag(paragraph.style.name)}{paragraph.text}{self._get_closing_tag(paragraph.style.name)}\n"
        documents = [
            SourceDocument(
                content=output,
                source=document_url,
                offset=0,
                page_number=0,
                global_business=metadata["global_business"],
                divisions_and_areas=metadata["divisions_and_areas"],
                tags=', '.join(metadata["tags"]),
                region=metadata["region"],
                country=metadata["country"],
                language=metadata["language"],
                year=metadata["year"],
                period=metadata["period"],
                importance=metadata["importance"],
                security=metadata["security"],
                origin=metadata["origin"],
                domain=metadata["domain"]
            )
        ]
        return documents
