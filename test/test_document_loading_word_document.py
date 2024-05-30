import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from backend.utilities.document_loading.WordDocument import WordDocumentLoading

@pytest.fixture
def word_document_loader():
    return WordDocumentLoading()

def create_fake_docx():
    file_stream = BytesIO()
    document = Document()
    document.add_paragraph("Test paragraph")
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def test_load_document(word_document_loader):
    # Mock the requests.get response with valid DOCX content
    with patch('requests.get') as mocked_get:
        mocked_get.return_value.content = create_fake_docx()
        
        # Mock the Document class from docx module
        with patch('docx.Document') as mocked_document:
            mocked_paragraph = MagicMock()
            mocked_paragraph.text = "Test paragraph"
            mocked_document.return_value.paragraphs = [mocked_paragraph]
            
            # Define the metadata
            metadata = {
                "global_business": "Tech",
                "divisions_and_areas": "R&D",
                "tags": ["innovation", "development"],
                "region": "Europe",
                "country": "Germany",
                "language": "English",
                "year": "2023",
                "period": "Q1",
                "importance": "High",
                "security": "Confidential",
                "origin": "Internal",
                "domain": "Engineering"
            }
            
            # Call the load method
            documents = word_document_loader.load("http://fakeurl.com/document.docx", metadata)
            
            # Assertions to verify the behavior
            assert len(documents) == 1
            assert documents[0].content == "<>Test paragraph</>\n"
            assert documents[0].source == "http://fakeurl.com/document.docx"
            assert documents[0].global_business == "Tech"
            assert documents[0].tags == "innovation, development"
            assert documents[0].language == "English"